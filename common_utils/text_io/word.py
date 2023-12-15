# encoding=utf-8
# created @2023/12/15
# created by zhanzq
#

import os

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


class Docx:
    @staticmethod
    def get_font_sz_map():
        font_sz_mp = {
            "初号": 42,
            "小初": 36,
            "一号": 26,
            "小一": 24,
            "二号": 22,
            "小二": 18,
            "三号": 16,
            "小三": 15,
            "四号": 14,
            "小四": 12,
            "五号": 10.5,
            "小五": 9,
            "六号": 7.5,
            "小六": 6.5,
            "七号": 5.5,
            "八号": 5,
        }

        return font_sz_mp

    def __init__(self, path):
        self.path = path
        if not os.path.exists(path):
            self.doc = Document()
        else:
            self.doc = Document(path)

        self.font_sz_map = self.get_font_sz_map()

    def set_font(self, run, **args):
        font_size = args.get("size", 20)
        font_name = args.get("name", "华文楷体")

        if type(font_size) is str:
            font_size = self.font_sz_map.get(font_size, 20)

        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.bold = args.get("bold", False)
        run.italic = args.get("italic", False)

        return

    def set_page(self, margin=(3.18, 3.18, 2.54, 2.54), column=1, column_space=0.5):
        section = self.doc.sections[-1]
        width, height = 21, 29.7  # A4
        l, r, t, b = margin
        section.page_width = Cm(width)
        section.page_height = Cm(height)

        section.left_margin = Cm(l)
        section.right_margin = Cm(r)
        section.top_margin = Cm(t)
        section.bottom_margin = Cm(b)

        if column > 1:
            cur_column = section._sectPr.xpath('./w:cols')[0].values()[-1]
            if cur_column != column:
                section = self.doc.add_section()

            cols = section._sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), '2')

        return

    def add_new_page(self, ):
        # 添加一个分页
        self.doc.add_page_break()

        return

    def add_figure(self, fig_path, width, height):
        self.doc.add_picture(fig_path, width=Cm(width), height=Cm(height))

        return

    def insert_table(self, data_lst):
        """
        向docx文档中插入表格
        :param data_lst: type(list), 表格数据
        """
        row, col = len(data_lst), len(data_lst[0])
        table = self.doc.add_table(rows=row, cols=col)

        for r in range(row):
            cells = table.rows[r].cells
            for c in range(col):
                cells[c].text = str(data_lst[r][c])

        return

    def extract_table(self, table_idx=0):
        t0 = self.doc.tables[table_idx]
        data_lst = []
        for i in range(len(t0.rows)):
            row = []
            for j in range(len(t0.columns)):
                row.append(t0.cell(i, j).text)
            data_lst.append(row)

        return data_lst

    def save(self, ):
        self.doc.save(self.path)
        print(f"{self.path} 文档已保存")
        return


if __name__ == "__main__":
    doc = Docx(path="/Users/zhanzq/Downloads/test.docx")
    doc.insert_table([["hello", "world"]])
    doc.save()
